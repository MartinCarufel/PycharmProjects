# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.



t_head = "Solution 1"

def read_xl_to_table(file):
    import openpyxl

    xtable = []

    xl = openpyxl.load_workbook(file)
    xl1 = xl.active
    for row in range(0, xl1.max_row):
        r = []
        for col in xl1.iter_cols(1, xl1.max_column):
            r.append(col[row].value)
        xtable.append(r)
    return xtable

def print_table(table):
    for row in table:
        for col in row:
            print(col, end='    ')
        print()

def word1():

    """
    Pour utiliser un style particulier, ouvrir un fichier Word, applique le(s) style
(s) et sauegarder le documenavec le style. Ensuite faire un delete des élments et sauvegarder a nouveau,
    Lors de l'instance Dcument() mettre le document vide créé a ce moment le doc vide contient les style particuliers.

    """

    from docx.document import Document

    try:
        d = Document("w_template.docx")
    except TypeError:
        from docx import Document
        d = Document("w_template.docx")

    xtb = read_xl_to_table("Data_summary.xlsx")

    d.add_heading(t_head, 1)
    d.add_paragraph()
    d.add_paragraph()
    table_width = len(xtb[0])
    w_table = d.add_table(1, table_width)
    for row_idx in range(0, len(xtb)):
        if row_idx == 0:
            xl_row = w_table.rows[row_idx].cells
        else:
            xl_row = w_table.add_row().cells
        data_row = xtb[row_idx]
        for col_idx in range(0, table_width):
            xl_row[col_idx].text = str(data_row[col_idx])

    # w_table.style = 'Plain Table 1'
    w_table.style = 'report table'
    d.save("test.docx")



# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    word1()
    # print_table(read_xl_to_table("Data_summary.xlsx"))


# See PyCharm help at https://www.jetbrains.com/help/pycharm/

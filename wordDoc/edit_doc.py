
from docx.document import Document

try:
    d = Document("2020-05-11 - Martin Carufel.docx")
except TypeError:
    from docx import Document
    d = Document("2020-05-11 - Martin Carufel.docx")

# Insert paragraph at specific position
d.paragraphs[3].insert_paragraph_before("Salut")

# Add in current paragraph
d.paragraphs[2].add_run("TUTU, TOTO")

print(d.tables)     # Get table list
print(d.tables[0].cell(0, 0).text)    #get cell containt

d.tables[0].cell(0, 0).text = "tata"       #Change cell
print(d.tables[0].cell(0, 0).text)
print(len(d.tables[0].rows))    #get length in row of table
d.tables[0].add_row()
d.tables[0].cell(4, 0).text = "Last L, col 1"
d.save("mod1.docx")


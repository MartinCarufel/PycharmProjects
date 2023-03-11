from shareplum import Site
from requests_ntlm import HttpNtlmAuth
import io
import urllib.request
from docx import Document
from docx.shared import Inches

import base64
def create_onedrive_directdownload (onedrive_link):
    data_bytes64 = base64.b64encode(bytes(onedrive_link, 'utf-8'))
    data_bytes64_String = data_bytes64.decode('utf-8').replace('/','_').replace('+','-').rstrip("=")
    resultUrl = f"https://api.onedrive.com/v1.0/shares/u!{data_bytes64_String}/root/content"
    return resultUrl

document = Document()
p = document.add_paragraph()
r = p.add_run()
# url = r"https://uploads-ssl.webflow.com/6014e375fb2fcc14f5e89e9b/6182f2bb0222842cde8a264e_spirit-plane-p-1600.png"
url = r"https://stgcs.sharepoint.com/:i:/r/sites/CA02_Proj_DW_Product/Shared%20Documents/General/Product/IO%20Cart_all%20ios/DHF/100%20Verification%20%26%20Validation/101%20Verification/Hardware/HPC%203.5/Rev%201%20release/Test%20Run%201/T253181%20Test%20for%20IP30%20ingress%20protection/IMG_20221216_151705.jpg?csf=1&web=1&e=EDzPG8"
# url64 = create_onedrive_directdownload(url)
# print(url64)

# io_url = io.BytesIO(urllib.request.urlopen(url64).read())
# r.add_picture(io_url)
# document.save('demo.docx')






auth = HttpNtlmAuth('u120230@straumann.com', '130.4giro')
site = Site('https://stgcs-my.sharepoint.com/personal/u120230_straumann_com/', auth=auth)
print(site.list())
# data = sp_list.GetListItems('All Items', rowlimit=200)
# print(sp_list)
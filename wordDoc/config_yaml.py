import yaml
import requests
import json
from docx.document import Document
from docx.shared import Pt
from docx import Document
import logging
import datetime
import re

formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
handler = logging.FileHandler(filename='test.log', mode='w', encoding="UTF-8")
handler.setFormatter(formatter)
logger = logging.getLogger()
logger.setLevel(logging.DEBUG)
logger.addHandler(handler)

class Data_admin:

    def __init__(self, credential_user, credential_pwd):
        self.credential_user = credential_user
        self.credential_pwd = credential_pwd
        self.status_list = {}
        self.id_initial = {}

    def create_user_initial(self):
        """
        Fill dictionary with user id: report initial as per Quality process.
        """
        name_initial = {"Martin Carufel": "MCL",
                        "Deepa Nallappan": "DNN",
                        "Duke Vu": "DVU",
                        "Kavitha thirunavukkarasu": "KTU",
                        "Natalia Piworun": "NPN",
                        "Theo Legrais": "TLS",
                        "Vidian-Manuel Duraud": "VDD"}
        URL = "https://testrail.dwos.com/index.php?/api/v2/get_users"
        authentication = (self.credential_user, self.credential_pwd)
        head = {"Content-Type": "application/json"}
        r = requests.get(URL, auth=authentication, headers=head)
        for user_name, initial in name_initial.items():
            for ul in r.json():
                if ul["name"] == user_name:
                    self.id_initial[ul["id"]] = initial

    def create_status_list(self):
        """
        Fill class dictionary with corresponding status_id: report text
        """
        status_text = {1:"Pass",
                       2:"Blocked",
                       3:"Untested",
                       4:"Retest",
                       5:"Fail",
                       7:"Pass w/Dev"}

        URL = "https://testrail.dwos.com/index.php?/api/v2/get_statuses"
        authentication = (self.credential_user, self.credential_pwd)
        head = {"Content-Type": "application/json"}
        r = requests.get(URL, auth=authentication, headers=head)
        for id, status_text in status_text.items():
            for status in r.json():
                # print(type(status["id"]))
                if status["id"] == id:
                    self.status_list[id] = status_text


class Create_csv_doc():
    def __init__(self, credential_user, credential_pwd):
        self.credential_user = credential_user
        self.credential_pwd = credential_pwd
        self.config = self.read_yaml()
        self.testrail_parent_id = self.config["testrail parent section id"]
        self.docx_table_ref_id = {}
        if not self.config["use CSV template"]:
            self.docx_table_ref_id = self.config["table mapping"]
        self.child_list = []
        self.child_list2 = {}
        self.doc = Document(self.config["template path"])
        data_admin = Data_admin('martin.carufel@dental-wings.com', '18,Mac&Amo')
        data_admin.create_status_list()
        data_admin.create_user_initial()
        self.status_list = data_admin.status_list
        self.user_initial = data_admin.id_initial

    def read_yaml(self):
        """ A function to read YAML config file"""
        with open('config.yml') as f:
            config = yaml.safe_load(f)
        return config       # Return dict of the config file

    def get_child_section_id(self, parent_id):
        """
        parent_id: (int) no de la section parente
        Return: List of integer of Test Rail section_id that have matching parent id
        """
        URL = "https://testrail.dwos.com//index.php?/api/v2/get_sections/" + str(self.config["project id"])   # get as dict all section in the project
                                                            # (id, name, description, parent_id, suite_id ...
        r = requests.get(URL, auth=(self.credential_user, self.credential_pwd),
                         headers={"Content-Type": "application/json"})
        for i in r.json():
            if i["parent_id"] == parent_id:
                table_id = self.found_corresponding_section_id(i)
                self.docx_table_ref_id[i["id"]] = table_id
        logger.debug("Number of sub section found: {}".format(len(self.docx_table_ref_id)))

    def found_corresponding_section_id(self, t_case_json_object):
        """
        Method to identify by the section name where in word doc we should write
        :param t_case_json_object: is json element from a request.get_sections
        :return: Word table index to write data
        """
        table_id = None
        if t_case_json_object["name"].upper() == 'INSTALLATION QUALIFICATION':
            table_id = 2
        if t_case_json_object["name"].upper() == 'OPERATIONAL QUALIFICATION':
            table_id = 3
        if t_case_json_object["name"].upper() == 'PERFORMANCE QUALIFICATION':
            table_id = 4
        if t_case_json_object["name"].upper() not in ['INSTALLATION QUALIFICATION', 'OPERATIONAL QUALIFICATION',
                                                      'PERFORMANCE QUALIFICATION']:
            logging.error("section '{}' named '{}' is not a valid name, check section title spelling"
                          .format(t_case_json_object["id"] ,t_case_json_object["name"]))
            table_id = None
        return table_id

    def change_table_font(self, table, font_size_pt):
        """Table: docx x table ex: d.table[0] where d in docx object"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(font_size_pt)

    def docx_spec_table_filling(self):
        section = None
        URL = "https://testrail.dwos.com//index.php?/api/v2/get_cases/" + str(self.config["project id"]) + "&section_id="
        authentication = (self.credential_user, self.credential_pwd)
        head = {"Content-Type": "application/json"}

        for section, value in self.docx_table_ref_id.items():
            logging.debug("Create table for section: {}" .format(section))
            logging.debug("writing in table: {}  name: {}" .format(value, section))
            r = requests.get(URL + str(section), auth=authentication, headers=head)
            if r.status_code == 200:
                for i in range(len(r.json())):
                    element = r.json()[i]
                    logging.debug("There is {} table in word doc" .format(len(self.doc.tables)))
                    logging.debug("Start write row {} for test case: {} in table {}".format(i, element["id"], value))
                    self.doc.tables[value].add_row()
                    # Write in word col 1 the step number (1,2,3,...)
                    self.doc.tables[value].cell(i+1, 0).text = str(i+1)
                    # Write in word col 2 the TestRail test case name and truncate the first 3 char that should be TS_, VS_
                    self.doc.tables[value].cell(i+1, 1).text = element['title'][3:] + "\n" + "Ref Test Rail: " + str(
                        element['id'])
                    # Write in word col 3 requirements if element custom_io_requirement is absent place alternative text
                    try:
                        self.doc.tables[value].cell(i + 1, 2).text = element['custom_io_requirement']
                    except TypeError:
                        self.doc.tables[value].cell(i + 1, 2).text = "Not defined"
                        logging.warning("Test Rail variable 'custom_io_requirement' not defined for {} - {}"
                                        .format(element["id"], element["title"]))
                    # Write in word col 4 TC all expected result of every TC step.
                    try:
                        self.doc.tables[value].cell(i + 1, 3).text = o.get_step_expected_result(element)
                    except TypeError:
                        self.doc.tables[value].cell(i + 1, 3).text = "Not Defined"
                        logging.warning("Test Rail variable 'custom_steps_separated' not defined for {} - {}"
                                        .format(element["id"], element["title"]))
                    # Write in word an aggregate of two element field into word col 5
                    try:
                        self.doc.tables[value].cell(i + 1, 4).text = element['custom_string_objective_evidence'] + " / " \
                                                                   + element['custom_test_objev']
                    except TypeError:
                        self.doc.tables[value].cell(i + 1, 4).text = "Not Defined"
                        logging.warning("Test Rail variable 'custom_string_objective_evidence' or 'custom_test_objev' not defined for {} - {}"
                                        .format(element["id"], element["title"]))
                    logging.debug("Finish write row {} for test case: {} in table {}".format(i, element["id"], value))
                # Format the table font to 9pts
                self.change_table_font(self.doc.tables[value], 9)


    def get_tc_list_from_section(self, section_id):
        tc_list = []
        URL = "https://testrail.dwos.com//index.php?/api/v2/get_cases/" + str(self.config["project id"]) + "&section_id="
        authentication = (self.credential_user, self.credential_pwd)
        head = {"Content-Type": "application/json"}
        r = requests.get(URL + str(section_id), auth=authentication, headers= head)
        logging.debug("Get test cases list for section id {} from Test Rail response: {}" .format(section_id ,r.status_code))
        for tc in range(len(r.json())):
            tc_list.append(r.json()[tc]["id"])
        return tc_list

    def extract_all_step_result(self, custom_step_result):
        """
        :param custom_step_result: Provide the content of the "get_results_for_case["custom_step_results"]" of
        a valid result
        :return:
        """
        regex_patern = '!\\[\\]\\(index\\.php\\?/attachments/get/\\d+\\)'
        steps_results = []
        logging.debug("Content of field 'custom_step_results: {}".format(custom_step_result))
        for actual_result_idx in range(len(custom_step_result)):
            if custom_step_result[actual_result_idx]["actual"] != "":
                steps_results.append("Step {}: ".format(actual_result_idx+1))
                filtered_string = re.subn(regex_patern, "", custom_step_result[actual_result_idx]["actual"])
                steps_results.append(filtered_string[0])
                # steps_results.append(custom_step_result[actual_result_idx]["actual"] + "\n")
        return "\n".join(steps_results)


    def docx_report_table_filling(self):
        URL = "https://testrail.dwos.com/index.php?/api/v2/get_results_for_case/"
        authentication = (self.credential_user, self.credential_pwd)
        head = {"Content-Type": "application/json"}
        for section, value in self.docx_table_ref_id.items():
            tc_list = self.get_tc_list_from_section(str(section))
            logging.debug("Following test case will exported to word {}".format(tc_list))
            logging.debug("Create table for section: {}" .format(section))
            logging.debug("writing in table: {}  name: {}" .format(value, section))
            for tc in range(len(tc_list)):
                self.doc.tables[value].add_row()
                logging.info("Write test case {}".format(tc_list[tc]))
                r = requests.get(URL + str(
                    self.config["test report run id"]) + "/" + str(tc_list[tc]), auth=authentication, headers=head)
                logging.debug("Request data: {}".format(r.json()))
                if r.status_code == 200 and r.json() != []:
                    for last_result_id in range(len(r.json())):
                        test_result_aggregate = []
                        logging.info("Read result index {} - id: {}, status id value is {}"
                                     .format(last_result_id, r.json()[last_result_id]["id"],
                                             r.json()[last_result_id]["status_id"]))
                        # if r.json()[last_result_id]["status_id"] != None:
                        if isinstance(r.json()[last_result_id]["status_id"], int):
                            logging.debug("Result accepted, status_id: {}".format(r.json()[last_result_id]["status_id"]))
                            # Write to word step id in col 1
                            self.doc.tables[value].cell(tc + 1, 0).text = str(tc + 1)  # Fill step number 1, 2, 3 ....
                            # Write to word test result by collecting all step actual result field in col 2
                            test_result_aggregate.append(f"Test case ID: {tc_list[tc]}")
                            test_result_aggregate.append("Test Result ID: {}\n".format(r.json()[last_result_id]["test_id"]))
                            test_result_aggregate.append(self.extract_all_step_result(r.json()[last_result_id]["custom_step_results"]))
                            # self.doc.tables[value].cell(tc + 1, 1).text =
                            # self.doc.tables[value].cell(tc + 1, 1).text =
                            self.doc.tables[value].cell(tc + 1, 1).text = "\n".join(test_result_aggregate)
                                # "actual result for {}".format(tc_list[tc])      # result text
                            logging.info("Test case status: {}"
                                         .format(self.status_list[r.json()[last_result_id]["status_id"]]))
                            self.doc.tables[value].cell(tc + 1, 2).text = self.status_list[r.json()[last_result_id]["status_id"]]  # Pass / Fail
                            logging.info("datetime stamp: {} - {}"
                                         .format(r.json()[last_result_id]["created_on"],
                                        datetime.datetime.fromtimestamp(r.json()[last_result_id]["created_on"])
                                                 .strftime("%Y-%m-%d")))
                            self.doc.tables[value].cell(tc + 1, 3).text = datetime.datetime.fromtimestamp(r.json()[last_result_id]["created_on"]).strftime("%Y-%m-%d")          # Date
                            self.doc.tables[value].cell(tc + 1, 4).text = "Ini"       # Initial
                            break
                        else:
                            logging.debug(
                                "Result skipped, status_id: {}".format(r.json()[last_result_id]["status_id"]))

                    # self.doc.tables[value].add_row()
                    # logging.debug("Finish write row {} for test case: {} in table {}".format(tc, element["id"], value))
                    # if tc < len(r.json()) - 1:  # avoid adding extra line at the end of table
                    #     self.doc.tables[value].add_row()
                    #     logging.debug("add row")

                else:
                    print("not test result")
                    logging.debug("No test result for test case: {}".format(tc_list[tc]))
                self.change_table_font(self.doc.tables[value], 9)

                    # for i in range(len(r.json())):
                    #     element = r.json()[i]
                    #     logging.debug("There is {} table in word doc" .format(len(self.doc.tables)))
                    #     logging.debug("Start write row {} for test case: {} in table {}".format(i, tc, value))
                    #
                    #     self.doc.tables[value].cell(i+1, 0).text = str(i+1)   # Fill step number 1, 2, 3 ....
                    #     self.doc.tables[value].cell(i+1, 1).text = "actual result"
                    #     self.doc.tables[value].cell(i + 1, 2).text = "Pass / fail"
                    #     self.doc.tables[value].cell(i + 1, 3).text = "date"
                    #     self.doc.tables[value].cell(i + 1, 4).text = "Ini"
                    #     logging.debug("Finish write row {} for test case: {} in table {}".format(i, element["id"], value))
                    #     if i < len(r.json())-1:   # avoid adding extra line at the end of table
                    #         self.doc.tables[value].add_row()
                    #         logging.debug("add row")
                    # self.change_table_font(self.doc.tables[value], 9)


    def get_step_expected_result(self, test_case):
        list_of_expected_result = []
        try:
            for i in test_case['custom_steps_separated']:
                if i['expected'] != "":
                    list_of_expected_result.append(i['expected'])
                    return "\n".join(list_of_expected_result)
        except TypeError:
            return "No expected result defined in step"

    def save_doc(self):
        self.doc.save(self.config["output doc name"])
        logging.debug("File saved {}".format(self.config["output doc name"]))





if __name__ == "__main__":
    data_admin = Data_admin('martin.carufel@dental-wings.com', '18,Mac&Amo')
    data_admin.create_user_initial()
    data_admin.create_status_list()
    o = Create_csv_doc('martin.carufel@dental-wings.com', '18,Mac&Amo')
    if o.config["test report"]:
        o.docx_report_table_filling()
    else:
        if o.config["use CSV template"]:
            o.get_child_section_id(o.testrail_parent_id)
        else:
            o.docx_spec_table_filling()
    o.save_doc()

    # print(o.get_tc_list_from_section(9517))


    # if o.config["use CSV template"]:
    #     o.get_child_section_id(o.testrail_parent_id)
    # if o.config["test report"]:
    #     pass
    # else:
    #     o.docx_spec_table_filling()



import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

emu_mm = 36000

def set_col_width(table_obj, col_idx, width_mm):
    col = table.columns
    for cel in col[col_idx].cells:
        cel.width = width_mm * emu_mm

def set_row_to_bold(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

def set_table_font_size(table, size_pts):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size_pts)

def set_table_cell_background(cell, color_hex_val):
    cell_xml_element = cell._tc
    cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color_hex_val)
    cell_properties.append(shade_obj)

wd = docx.Document()
sections = wd.sections
sections[0].left_margin = 18*emu_mm
sections[0].right_margin = 18*emu_mm

table = wd.add_table(rows=2, cols=5)
table.style = "Table Grid"

set_col_width(table, 0, 10)
set_col_width(table, 1, 125)
set_col_width(table, 2, 15)
set_col_width(table, 3, 25)
set_col_width(table, 4, 10)

row = table.rows
cell = row[0].cells

cell[0].text = "Step"
cell[1].text = "Result Description"
cell[2].text = "Result (Pass or Fail (P/F)"
cell[3].text = "Test Date (yyyy-mm-dd)"
cell[4].text = "Initial"
set_table_font_size(table, 9)
set_row_to_bold(row[0])
for cell in table.rows[0].cells:
    set_table_cell_background(cell, "c0c0c0")

wd.save("Word_table.docx")
import requests
import json

from docx.document import Document
from docx.shared import Pt
try:
    d = Document()
except TypeError:
    from docx import Document
    d = Document()


def get_step_expected_result(test_case):
    list_of_expected_result = []

    for i in test_case['custom_steps_separated']:
        if i['expected'] != "":
            list_of_expected_result.append(i['expected'])

    return "\n".join(list_of_expected_result)


iurl = 'https://the-one-api.dev/v2'
command = '/book'

cr2 = ('martin.carufel@dental-wings.com', '18,Mac&Amo')
URL = "https://testrail.dwos.com//index.php?/api/v2/get_case/42400"
# result = requests.get('https://the-one-api.dev/v2/book')
# result = requests.get(iurl+command)
# my_dic = json.loads(result.text)

# result = requests.get(url=URL, headers={"Content-Type": "application/json"}, auth=cred)
# print(result.json())
# print(result.text)
# print(my_dic['total'])


# session = requests.session()
# session.auth('martin.carufel@dental-wings.com', '18,Mac&Amo')
# session.get(URL)

# r = requests.get(URL, auth=cr2, headers={"Content-Type": "application/json"})
# print(r.json())
#
# URL = "https://testrail.dwos.com//index.php?/api/v2/get_section/11362"
# r = requests.get(URL, auth=cr2, headers={"Content-Type": "application/json"})
# print(r.json())

d.add_table(1, 5, style='Table Grid')
d.tables[0].cell(0, 0).text = "Step"
d.tables[0].cell(0, 1).text = "Description"
d.tables[0].cell(0, 2).text = "Requirements"
d.tables[0].cell(0, 3).text = "Expected Result"
d.tables[0].cell(0, 4).text = "Test Methode / Objective Evidence"
row_id = 1



URL = "https://testrail.dwos.com//index.php?/api/v2/get_cases/2&section_id=11362"
r = requests.get(URL, auth=cr2, headers={"Content-Type": "application/json"})
for i in r.json():
    d.tables[0].add_row()
    d.tables[0].cell(row_id, 0).text = str(row_id)
    d.tables[0].cell(row_id, 1).text = i['title'][3:] + "\n" + "Ref Test Rail: " + str(i['id'])
    d.tables[0].cell(row_id, 2).text = i['custom_io_requirement']
    d.tables[0].cell(row_id, 3).text = get_step_expected_result(i)
    d.tables[0].cell(row_id, 4).text = i['custom_string_objective_evidence'] + " / " + i['custom_test_objev']
    row_id += 1

for row in d.tables[0].rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(9)

d.save("template_filled.docx")

    # print(i)


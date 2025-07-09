# from spire.doc import *
# from spire.doc.common import *
import string

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.text.run import Run
import re
import csv


def get_doc_structure(doc):
    all_elements = []
    for child in doc._element.body.iterchildren():
        if isinstance(child, CT_P):
            all_elements.append(Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            all_elements.append(Table(child, doc))
    return all_elements

def get_table(file_structure, para_index):
    para_count = 0
    for idx, elem in enumerate(file_structure):
        if isinstance(elem, Paragraph):
            if para_count == para_index:
                # Scan ahead to find the next table
                for next_elem in file_structure[idx + 1:]:
                    if isinstance(next_elem, Table):
                        return next_elem
            para_count += 1

    return None  # No table found after the specified paragraph index

def get_test_case(document):
    """
    Get the test case paragraph id and the title
    :param document:
    :return: A list of tupple where the item [0] is the paragraph ID in word that the TC was found
    and the [1] is the test case title itself.
    """
    par = document.paragraphs
    test_case_list = []
    for i in range(len(par)):
        x = re.match("TC[0-9]{5}", par[i].text)
        if x != None:
            test_case_list.append((i, par[i].text.rstrip()))
            # print(i, end="- ")
            # print(par[i].text)
            pass
    return test_case_list

def clean_and_sort_req_list(input_list):
    return sorted(set(input_list))

def clean_and_split(input_str, target_list):
    """
    Splits a string by commas, removes non-printable characters like spaces and returns,
    and appends cleaned items to an existing list.

    Parameters:
    input_str (str): The input string containing comma-separated values.
    target_list (list): The list to which cleaned values will be appended.
    """
    # Remove unwanted characters and split
    cleaned_items = [
        ''.join(c for c in item if c in string.printable and not c.isspace())
        for item in input_str.split(',')
    ]
    target_list.extend(filter(None, cleaned_items))  # Avoid adding empty strings

def normalize_newlines(txt):
    """
    Replaces various new line and paragraph break indicators with '\\n'.
    """
    # Common replacements: carriage returns, paragraph symbols, HTML breaks, etc.
    return re.sub(r'(\r\n|\r|\n|¶| |<br\s*/?>)', '\\n', txt)

def break_new_line(o):
    result = []
    for para in o.paragraphs:
        text_parts = []
        for node in para._element.iter():
            if node.tag == qn('w:br'):
                text_parts.append("\\n")  # Word line break (Shift+Enter)
            elif node.tag == qn('w:t'):
                text_parts.append(node.text or "")
        para_text = ''.join(text_parts)
        result.append(para_text)

        # Only join with '\n' between paragraphs if the paragraph has actual text
    return '\\n'.join(result)

def extract_text_with_all_breaks_as_backslash_n(doc):

    full_text = []

    for para in doc.paragraphs:
        text_parts = []

        for node in para._element.iter():
            if node.tag == qn('w:br'):
                # print("new line")
                text_parts.append("\\n")  # Line break inside a paragraph
            elif node.tag == qn('w:t'):
                text_parts.append(node.text or "")

        para_text = ''.join(text_parts)
        # print(para_text)
        full_text.append(para_text)

    # Replace paragraph break with \n too
    # print(full_text)
    return "\\n".join(full_text)
def csv_construct_header(header_l):
    return ",".join(header_l) + "\n"
def csv_line_construct_tc(o):
    return f",\"{o[1]}\",,,\n"

def csv_construct_test_step(o):
    # print(o)
    x = extract_text_with_all_breaks_as_backslash_n(o)
    print(x)

    para_list = [x.text for x in o.paragraphs]
    # print(para_list)
    new_text = "\\n".join(para_list)

    clean_text = normalize_newlines(new_text)
    return f",,,\"{clean_text}\"\n"

def main():
    path_file = "Unit_test/test_data/DEV-0044600 STMN IOS Main Application Verification Specifications Rev 4.docx"
    document = Document(path_file)
    s = get_doc_structure(document)
    test_case_list = get_test_case(document)
    # print(test_case_list)

    with open("export.csv", mode="w", encoding="UTF-8") as f:
        f.write(csv_construct_header(["ID","Work Item Type","Title","Test Step","Step Action","Step Expected"]))
        for item in test_case_list:
            f.write(csv_line_construct_tc(item))
            next_table = get_table(s, item[0])
            """for row in next_table.rows:
                # f.write(csv_construct_test_step(row.cells[1]))
                # print(row.cells[1].paragraphs)
                f.write(csv_construct_test_step(row.cells[1]))
                # for p in row.cells[1].paragraphs:
                #     print(p.text)"""

            for i in range(len(next_table.rows)):
                # f.write(csv_construct_test_step(row.cells[1]))
                # print(row.cells[1].paragraphs)
                if i > 0:
                    f.write(csv_construct_test_step(next_table.rows[i].cells[1]))
                # for p in row.cells[1].paragraphs:
                #     print(p.text)


    """next_table = get_table(s, 9)
    for row in next_table.rows:
        print([cell.text for cell in row.cells])
    pass"""


if __name__ == "__main__":
    main()